import io
import os
import re
import sys
import tempfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from pathlib import Path

import aspose.words as aw
import debugpy
import language_tool_python
from colorama import Fore, Style
from docx import Document, table
from docx2python import docx2python
from docx.shared import Cm
from PySide6.QtCore import QSettings, Qt, QThread, Signal
from PySide6.QtGui import (
    QAction,
    QColor,
    QFont,
    QTextBlock,
    QTextCharFormat,
    QTextCursor,
    QTextDocument,
)
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QColorDialog,
    QComboBox,
    QDialog,
    QFileDialog,
    QInputDialog,
    QLabel,
    QMainWindow,
    QPushButton,
    QSplitter,
    QStyle,
    QTextEdit,
    QVBoxLayout,
)

templates = [
    {
        "name": "Smartcat",
        "simple": False,
        "row": 0,
        "source_col_index": 1,
        "target_col_index": 2,
    },
    {
        "name": "MemoQ",
        "simple": False,
        "row": 0,
        "source_col_index": 1,
        "target_col_index": 2,
    },
    {
        "name": "Memsouce",
        "simple": False,
        "row": 0,
        "source_col_index": 3,
        "target_col_index": 4,
    },
    {
        "name": "Target General",
        "simple": True,
    },
]

error_type_color_map = {
    "uncategorized": Qt.GlobalColor.magenta,
    "misspelling": Qt.GlobalColor.red,
    "style": Qt.GlobalColor.blue,
}


class TextDisplay(QTextEdit):
    def __init__(self):
        super().__init__()

    def mouseMoveEvent(self, e):
        self.anchor = self.anchorAt(e.position().toPoint())
        if self.anchor:
            QApplication.setOverrideCursor(Qt.CursorShape.PointingHandCursor)
        else:
            QApplication.setOverrideCursor(Qt.CursorShape.ArrowCursor)

    def mousePressEvent(self, e):
        self.anchor = self.anchorAt(e.position().toPoint())

        if self.anchor:
            print(f"Anchor: {self.anchor}")

        # if self.anchor:
        #     QApplication.setOverrideCursor(Qt.CursorShape.PointingHandCursor)

    def mouseReleaseEvent(self, e):
        # if self.anchor:
        # QDesktopServices.openUrl(QUrl(self.anchor))
        # QApplication.setOverrideCursor(Qt.CursorShape.ArrowCursor)
        # self.anchor = None
        pass


class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.recentFiles: list[str] = []
        self.initUI()

    def initUI(self):
        self.splitter = QSplitter()
        self.splitter.setOrientation(Qt.Orientation.Horizontal)
        self.setCentralWidget(self.splitter)

        self.textDisplay = TextDisplay()
        self.textDisplay.setAcceptRichText(True)
        self.textDisplay.setReadOnly(True)
        self.textDisplay.setStyleSheet("background-color: white;" "color: black;")
        self.splitter.addWidget(self.textDisplay)

        self.errorDisplay = QTextEdit()
        self.errorDisplay.setAcceptRichText(True)
        self.errorDisplay.setReadOnly(True)
        self.splitter.addWidget(self.errorDisplay)

        self.errorDisplay.setText("Errors will be displayed here.")

        cursor = QTextCursor(self.errorDisplay.document())
        cursor.setPosition(0)
        cursor.insertHtml("<b>")
        # cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, 5)
        cursor.setPosition(5)
        cursor.insertHtml("</b>")

        # Load recent files from QSettings
        recentFiles = QSettings().value("recentFiles", [])

        # Check if this is a list of strings or a single string
        if isinstance(recentFiles, str):
            recentFiles = [recentFiles]

        self.recentFiles = recentFiles

        openAction = QAction("Open", self)
        openAction.setShortcut("Ctrl+O")
        openAction.setStatusTip("Open File")
        openAction.setIcon(
            self.style().standardIcon(QStyle.StandardPixmap.SP_DialogOpenButton)
        )
        openAction.triggered.connect(self.openFile)

        exitAction = QAction("Exit", self)
        exitAction.triggered.connect(self.close)

        menubar = self.menuBar()
        fileMenu = menubar.addMenu("File")
        fileMenu.addAction(openAction)

        self.recentMenu = fileMenu.addMenu("Recent Files")
        self.updateRecentFilesMenu()
        
        preferencesAction = QAction("Preferences", self)
        preferencesAction.triggered.connect(self.openPreferences)

        fileMenu.addAction(preferencesAction)
        fileMenu.addSeparator()
        fileMenu.addAction(exitAction)

        self.setWindowTitle("pyLanguageTool")

        self.statusBar()

        # Add an icon bar
        self.toolbar = self.addToolBar("Toolbar")
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, self.toolbar)
        self.toolbar.addAction(openAction)

        # Add template dropdown
        self.templateComboBox = QComboBox()
        self.templateComboBox.addItems([template["name"] for template in templates])
        self.templateComboBox.currentIndexChanged.connect(self.templateChanged)
        self.toolbar.addWidget(self.templateComboBox)

        # Add checkbox "Remove tags"
        self.removeTagsCheckBox = QCheckBox("Remove tags")
        self.toolbar.addWidget(self.removeTagsCheckBox)

        # Maximize the window
        # self.showMaximized()
        self.resize(1200, 800)

        self.loadWindowPosition()

        self.show()

        self.language_tool = language_tool_python.LanguageTool("de-DE")

        self.errors: dict = {}

        self.current_template = templates[0]

    def openPreferences(self):
        preferencesWindow = PreferencesWindow(self)
        preferencesWindow.exec()

    def templateChanged(self, index):
        template_name = self.templateComboBox.currentText()
        template = next(
            (template for template in templates if template["name"] == template_name),
            None,
        )
        if template:
            self.current_template = template

    def closeEvent(self, event):
        self.saveWindowPosition()

        QSettings().setValue("recentFiles", self.recentFiles)

        event.accept()

    def read_docx_tables(self, file_path):
        """
        Reads all the tables in a docx file and returns them as a list.
        """
        document = Document(file_path)
        tables = document.tables

        return tables

    def extract_table_columns(self, table, columns, num_rows=-1):
        """
        Extracts the contents of the specified columns of a table up to a specified number of rows.
        """
        extracted_columns = [[] for _ in range(len(columns))]

        if num_rows == -1:
            num_rows = len(table.rows)

        for i, row in enumerate(table.rows):
            if num_rows > -1 and i >= num_rows:
                break
            for j, col in enumerate(columns):
                if col < len(row.cells):
                    extracted_columns[j].append(row.cells[col].text.strip())
            if i % 100 == 0:
                print(
                    f"{Fore.GREEN}Current row number: {i} of {len(table.rows)}{Style.RESET_ALL}, First column: {Fore.BLUE}{extracted_columns[0][-1]}{Style.RESET_ALL}"
                )
        return extracted_columns

    def printError(self, cursor: QTextCursor, error: dict):
        # Get color based on error type
        error_type = error["Error"].split(" - ")[0]
        color = QColor(error_type_color_map.get(error_type, Qt.GlobalColor.yellow))

        for field_name, field_value in error.items():
            if field_name == "Offset" or field_name == "Length":
                continue
            # Ignore replacements if there are none
            if field_name == "Replacements" and field_value != [" "]:
                continue

            block_format = cursor.blockFormat()
            # block_format.setTopMargin(10)
            block_format.setBackground(color.lighter(190))
            # block_format.setForeground(Qt.GlobalColor.black)

            cursor.setBlockFormat(block_format)
            cursor.insertBlock()
            format = QTextCharFormat()

            format.setForeground(Qt.GlobalColor.black)
            format.setFontWeight(QFont.Weight.Bold)

            cursor.insertText(f"{field_name}: ", format)

            format.setFontWeight(QFont.Weight.Normal)
            # format = QTextCharFormat()
            # format.setForeground(Qt.GlobalColor.black)

            cursor.insertText(f"{field_value}\n", format)

    def fileLoaded(self, text: str):
        cursor = QTextCursor(self.errorDisplay.document())

        for error in self.errors.values():
            self.printError(cursor, error)

        block_format = cursor.blockFormat()
        block_format.setBackground(QColor(Qt.GlobalColor.white))
        cursor.insertBlock()
        cursor.insertText("\n")

        formatted_text = self.formatText(text)
        self.textDisplay.setDocument(formatted_text)

        if self.removeTagsCheckBox.isChecked():
            text = re.sub(r"<.*?>", "", text)

        self.addRecentFile(self.fileLoaderWorker.file_name)
        self.statusBar().showMessage("File loaded")

    def openFile(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "Open File")
        if file_name:
            self.fileLoaderWorker = FileLoaderWorker(self, file_name)
            self.fileLoaderWorker.fileLoaded.connect(self.fileLoaded)
            self.fileLoaderWorker.start()

    def formatText(self, text):
        document = QTextDocument()
        cursor = QTextCursor(document)

        # Display the text in the QTextEdit
        # cursor.insertText(text)

        format = QTextCharFormat()

        # Get a list of all keys (offsets) in the errors dictionary
        keys = list(self.errors.keys())

        current_error = None

        # Loop through the text and add formatting based on the offset and length
        for character in text:
            offset = cursor.position()
            if offset in keys:
                error = self.errors[offset]
                format.setFontUnderline(True)

                error_type = error["Error"].split(" - ")[0]

                format.setUnderlineColor(
                    error_type_color_map.get(error_type, Qt.GlobalColor.black)
                )
                format.setUnderlineStyle(
                    QTextCharFormat.UnderlineStyle.SpellCheckUnderline
                )
                format.setToolTip(
                    f"{error['Error']}\n{error['Message']}\n→ {error['Replacements']}\nContext: {error['Context']}\n{error['Sentence']}"
                )
                format.setAnchor(True)
                format.setAnchorHref(f"#{offset}")
                current_error = error
            elif current_error:
                if offset == error["Offset"] + error["Length"]:
                    format = QTextCharFormat()
                    current_error = None

            cursor.insertText(character, format)
        return document

    def addRecentFile(self, file_name):
        if file_name in self.recentFiles:
            self.recentFiles.remove(file_name)
        self.recentFiles.insert(0, file_name)
        self.updateRecentFilesMenu()

    def openRecentFile(self, file_name):
        if file_name:
            self.fileLoaderWorker = FileLoaderWorker(self, file_name)
            self.fileLoaderWorker.fileLoaded.connect(self.fileLoaded)
            self.fileLoaderWorker.start()
        self.addRecentFile(file_name)

    def updateRecentFilesMenu(self):
        self.recentMenu.clear()
        for i, fileName in enumerate(self.recentFiles):
            action = QAction(f"{i+1}: {fileName}", self)

            action.triggered.connect(
                lambda _, fileName=fileName: self.openRecentFile(fileName)
            )
            self.recentMenu.addAction(action)

    def saveWindowPosition(self):
        settings = QSettings()
        settings.setValue("window/position", self.pos())
        settings.setValue("window/size", self.size())
        settings.setValue("recentFiles", self.recentFiles)

    def loadWindowPosition(self):
        settings = QSettings()
        pos = settings.value("window/position", self.pos())
        size = settings.value("window/size", self.size())
        self.move(pos)
        self.resize(size)


class FileLoaderWorker(QThread):
    fileLoaded = Signal(str)

    def __init__(self, text_editor: TextEditor, file_name: str):
        super().__init__()
        self.text_editor = text_editor
        self.file_name = file_name

    def run(self):
        # debugpy.debug_this_thread()
        self.text_editor.statusBar().showMessage(f"Loading {self.file_name}...")
        with open(self.file_name, "r") as file:
            extension = Path(self.file_name).suffix.lstrip(".")
            match extension:
                case "docx" | "doc" | "rtf":
                    file_path = None

                    match extension:
                        case "rtf":
                            # Load file as bytesio
                            with open(self.file_name, "rb") as f:
                                data = f.read()

                            stream = io.BytesIO(data)
                            doc = aw.Document(stream)

                            # Save as docx
                            stream = io.BytesIO()
                            doc.save(stream, aw.SaveFormat.DOCX)
                            stream.seek(0)

                            file_path = stream
                        case "docx" | "doc":
                            file_path = self.file_name

                    current_template = self.text_editor.current_template

                    if current_template.get("simple", True):
                        with docx2python(file_path) as docx_content:
                            text = docx_content.text
                    else:
                        tables = self.text_editor.read_docx_tables(file_path)

                        text = ""

                        table = tables[current_template["row"]]
                        column_index = current_template["target_col_index"]

                        target1 = self.text_editor.extract_table_columns(
                            table, [column_index]
                        )[0]

                        text = "\n".join(target1)
                case "xliff":
                    # Parse the xliff file
                    root = ET.parse(self.file_name).getroot()

                    # source = root.findall(".//{urn:oasis:names:tc:xliff:document:1.2}source")
                    target = root.findall(
                        ".//{urn:oasis:names:tc:xliff:document:1.2}target"
                    )

                    text = ""

                    for t in target:
                        if t.text:
                            text += t.text + "\n"

                case _:
                    text = file.read()

            self.text_editor.statusBar().showMessage(
                f"Checking {self.file_name} with LanguageTool..."
            )
            self.checkSegment(text)

        self.fileLoaded.emit(text)

    def select_column(self, column_names) -> int:
        column_name, ok = QInputDialog.getItem(
            self.text_editor,
            "Select Column",
            "Select the column to check:",
            column_names,
            0,
            False,
        )

        column_index = column_names.index(column_name)
        return column_index

    def select_table(self, table_names, tables) -> int:
        if len(tables) > 1:
            table_name, ok = QInputDialog.getItem(
                self.text_editor,
                "Select Table",
                "Select the table to check:",
                table_names,
                0,
                False,
            )
            if not ok:
                table_name = table_names[0]
        else:
            table_name = table_names[0]

        table_index = table_names.index(table_name) + 1
        return table_index

    def checkSegment(self, segment: str):
        matches = self.text_editor.language_tool.check(segment)
        for match in matches:
            print(f"{Fore.RED}Error: {match.message}{Style.RESET_ALL}")
            error_type = f"{match.ruleIssueType} - {match.category}"
            error = {
                "Error": error_type,
                "Message": match.message,
                "Replacements": match.replacements,
                "Context": match.context,
                "Sentence": match.sentence,
                "Offset": match.offset,
                "Length": match.errorLength,
            }
            self.text_editor.errors[match.offset] = error


class PreferencesWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Preferences")
        self.layout = QVBoxLayout()

        self.errorColors = {}

        for error_type, color in error_type_color_map.items():
            label = QLabel(f"{error_type}:")
            self.layout.addWidget(label)

            colorButton = QPushButton()
            colorButton.setStyleSheet(f"background-color: {color.name}")
            colorButton.clicked.connect(
                lambda _, error_type=error_type: self.setColor(error_type)
            )
            self.layout.addWidget(colorButton)

            self.errorColors[error_type] = colorButton

        self.setLayout(self.layout)
        self.setMinimumWidth(300)  # Set minimum width for the dialog

    def setColor(self, error_type):
        color = QColorDialog.getColor()
        if color.isValid():
            self.errorColors[error_type].setStyleSheet(
                f"background-color: {color.name}"
            )
            error_type_color_map[error_type] = color


if __name__ == "__main__":
    # Set the organization name and application name for QSettings
    QApplication.setOrganizationName("Andre Jonas")
    QApplication.setApplicationName("pyLanguageTool")

    # Set the user config directory
    config_dir = QSettings().fileName()
    config_dir = os.path.dirname(config_dir)
    QSettings.setPath(QSettings.Format.IniFormat, QSettings.Scope.UserScope, config_dir)

    app = QApplication(sys.argv)
    editor = TextEditor()
    sys.exit(app.exec())
