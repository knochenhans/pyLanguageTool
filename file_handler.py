import io
from pathlib import Path
from docx import Document
from docx2python import docx2python
from colorama import Fore, Style

import aspose.words as aw
import xml.etree.ElementTree as ET


class FileHandler:
    def __init__(self, text_editor):
        self.text_editor = text_editor

    def read_docx_tables(self, file_path):
        """
        Reads all the tables in a docx file and returns them as a list.
        """
        document = Document(file_path)
        tables = document.tables

        return tables

    def extract_table_columns(self, table, columns, num_rows=-1):
        """
        Extracts the contents of the specified columns of a table up to a specified number of rows.
        """
        extracted_columns = [[] for _ in range(len(columns))]

        if num_rows == -1:
            num_rows = len(table.rows)

        for i, row in enumerate(table.rows):
            if num_rows > -1 and i >= num_rows:
                break
            for j, col in enumerate(columns):
                if col < len(row.cells):
                    extracted_columns[j].append(row.cells[col].text.strip())
            if i % 100 == 0:
                print(
                    f"{Fore.GREEN}Current row number: {i} of {len(table.rows)}{Style.RESET_ALL}, First column: {Fore.BLUE}{extracted_columns[0][-1]}{Style.RESET_ALL}"
                )
        return extracted_columns

    def load_file(self, file_name):
        with open(file_name, "r") as file:
            extension = Path(file_name).suffix.lstrip(".")
            match extension:
                case "docx" | "doc" | "rtf":
                    file_path = None

                    match extension:
                        case "rtf":
                            # Load file as bytesio
                            with open(file_name, "rb") as f:
                                data = f.read()

                            stream = io.BytesIO(data)
                            doc = aw.Document(stream)

                            # Save as docx
                            stream = io.BytesIO()
                            doc.save(stream, aw.SaveFormat.DOCX)
                            stream.seek(0)

                            file_path = stream
                        case "docx" | "doc":
                            file_path = file_name

                    current_template = self.text_editor.current_template

                    if current_template.get("simple", True):
                        with docx2python(file_path) as docx_content:
                            text = docx_content.text
                    else:
                        tables = self.read_docx_tables(file_path)

                        text = ""

                        table = tables[current_template["row"]]
                        column_index = current_template["target_col_index"]

                        target1 = self.extract_table_columns(table, [column_index])[0]

                        text = "\n".join(target1)
                case "xliff":
                    # Parse the xliff file
                    root = ET.parse(file_name).getroot()

                    # source = root.findall(".//{urn:oasis:names:tc:xliff:document:1.2}source")
                    target = root.findall(
                        ".//{urn:oasis:names:tc:xliff:document:1.2}target"
                    )

                    text = ""

                    for t in target:
                        if t.text:
                            text += t.text + "\n"

                case "mxliff":
                    # Parse the xliff file
                    root = ET.parse(file_name).getroot()

                    # source = root.findall(".//{urn:oasis:names:tc:xliff:document:1.2}source")
                    target = root.findall(
                        ".//{urn:oasis:names:tc:xliff:document:1.2}target"
                    )

                    text = ""

                    for t in target:
                        if t.text:
                            text += t.text + "\n"

                case _:
                    text = file.read()

        return text
